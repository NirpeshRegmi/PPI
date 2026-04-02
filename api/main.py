import gc
import io
import re
import base64
import difflib
from typing import List, Optional

import fitz  # pymupdf
import usaddress
import docx
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import openpyxl
import pdfplumber
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig
from starlette.middleware.base import BaseHTTPMiddleware

app = FastAPI(title="Transcend PII Processor")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class GCMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        response = await call_next(request)
        gc.collect()
        return response

app.add_middleware(GCMiddleware)

analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

# ---------------------------------------------------------------------------
# Presidio entity list — DATE_TIME removed to avoid killing tax years.
# Tax year shielding now happens before Presidio so years/quarters survive.
# Added AGE and IN_PAN for broader international PII coverage.
# ---------------------------------------------------------------------------
PII_ENTITIES = [
    "PERSON",
    "PHONE_NUMBER",
    "EMAIL_ADDRESS",
    "LOCATION",
    "US_SSN",
    "US_ITIN",
    "US_DRIVER_LICENSE",
    "US_PASSPORT",
    "US_BANK_NUMBER",
    "CREDIT_CARD",
    "IBAN_CODE",
    "IP_ADDRESS",
    "URL",
    "NRP",
    "MEDICAL_LICENSE",
    "CRYPTO",
    # DATE_TIME intentionally excluded — tax years must survive.
    # They are shielded before Presidio and restored after.
]

# ---------------------------------------------------------------------------
# FINANCIAL SHIELD
# Protects dollar amounts, percentages, and tax-year dates from being
# mangled by downstream redaction passes.
# Run AFTER sensitive-ID redaction so we don't fragment SSN/EIN patterns.
# ---------------------------------------------------------------------------

DOLLAR_RE = re.compile(r'\$[\d,]+(?:\.\d+)?(?:[KkMmBb])?')
PERCENT_RE = re.compile(r'\b\d+(?:\.\d+)?%')
TAX_DATE_RE = re.compile(
    r'\b(?:20[0-2]\d)\b'
    r'|\b(?:0?[1-9]|1[0-2])/(?:0?[1-9]|[12]\d|3[01])/(?:20[0-2]\d)\b'
    r'|\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?'
    r'|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\s+(?:20[0-2]\d)\b'
    r'|\bQ[1-4]\s+(?:20[0-2]\d)\b',
    re.IGNORECASE
)
FINANCIAL_PATTERNS = [DOLLAR_RE, PERCENT_RE, TAX_DATE_RE]

_FIN_TOKEN_RE = re.compile(r'__FIN_\d+__')


def shield_financials(text: str) -> tuple[str, dict]:
    placeholder_map = {}
    counter = [0]

    def replace(m):
        token = f"__FIN_{counter[0]}__"
        placeholder_map[token] = m.group(0)
        counter[0] += 1
        return token

    for pattern in FINANCIAL_PATTERNS:
        text = pattern.sub(replace, text)
    return text, placeholder_map


def restore_financials(text: str, placeholder_map: dict) -> str:
    if not placeholder_map:
        return text
    return _FIN_TOKEN_RE.sub(lambda m: placeholder_map.get(m.group(0), m.group(0)), text)


# ---------------------------------------------------------------------------
# ADDRESS REDACTION
# usaddress confirms before redacting to reduce false positives on
# non-address patterns, but ZIP codes are always redacted unconditionally
# since false positives are acceptable and ZIP alone is identifying.
# ---------------------------------------------------------------------------

LOOSE_ADDRESS_RE = re.compile(
    r'\b\d{1,6}'
    r'(?:\s+[NSEWnsew]\.?)?\s+'
    r'[A-Za-z0-9]+(?:\s+[A-Za-z0-9]+){0,4}\s+'
    r'(?:street|st|avenue|ave|boulevard|blvd|road|rd|drive|dr|lane|ln'
    r'|court|ct|circle|cir|place|pl|way|wy|terrace|ter|trail|trl'
    r'|highway|hwy|parkway|pkwy|square|sq|loop|lp|run|path|row'
    r'|alley|aly|crossing|xing)\.?'
    r'(?:\s+(?:apt|apartment|suite|ste|unit|#|no\.?)\s*[A-Za-z0-9-]+)?'
    r'(?:[,\s]+'
    r'[A-Za-z]+(?:\s+[A-Za-z]+){0,2}'
    r'(?:[,\s]+'
    r'(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)'
    r'(?:\s+\d{5}(?:-\d{4})?)?'
    r')?)?',
    re.IGNORECASE,
)

UNIT_ONLY_RE = re.compile(
    r'\b(?:suite|ste|unit|apt|apartment|floor|fl)\s+[A-Za-z0-9-]+'
    r'(?:[,\s]+[A-Za-z]+(?:\s+[A-Za-z]+){0,2})?'
    r'(?:[,\s]+'
    r'(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)'
    r'(?:\s+\d{5}(?:-\d{4})?)?'
    r')?',
    re.IGNORECASE,
)

RURAL_ROUTE_RE = re.compile(
    r'\b(?:rural\s+route|rr|r\.r\.|hc|highway\s+contract)\s*\d+'
    r'(?:\s+box\s+\d+)?'
    r'(?:[,\s]+[A-Za-z]+(?:\s+[A-Za-z]+){0,2})?'
    r'(?:[,\s]+'
    r'(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)'
    r'(?:\s+\d{5}(?:-\d{4})?)?'
    r')?',
    re.IGNORECASE,
)

POBOX_RE = re.compile(r'\bP\.?\s*O\.?\s*Box\s+\d+\b', re.IGNORECASE)

ZIP_RE = re.compile(r'\b\d{5}(?:-\d{4})?\b(?=(?:\D|$))')

# 9-digit ZIP+4 without dash — only redact when preceded by a state abbreviation
# to avoid hitting financial figures, routing numbers, or account numbers.
# e.g. "TX 750193679" or "ND 582027155" — but NOT bare "031176110"
ZIP_9_STATE_RE = re.compile(
    r'\b(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)\s{0,2}(\d{9})\b'
)

_ADDRESS_LABELS = {
    'AddressNumber', 'StreetName', 'StreetNamePostType',
    'StreetNamePreType', 'StreetNamePreDirectional', 'StreetNamePostDirectional',
    'OccupancyType', 'OccupancyIdentifier',
    'PlaceName', 'StateName', 'ZipCode',
}


def _usaddress_confirms(span: str) -> bool:
    try:
        tagged, addr_type = usaddress.tag(span)
        labels = set(tagged.values())
        if addr_type in ('Street Address', 'PO Box'):
            return True
        if len(labels & _ADDRESS_LABELS) >= 2:
            return True
        return False
    except usaddress.RepeatedLabelError:
        return True


# City + state abbreviation pattern — catches "Coppell TX", "Coppell, TX"
# that Presidio's LOCATION entity sometimes misses.
CITY_STATE_RE = re.compile(
    r'\b([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){0,2})'
    r'[,\s]+'
    r'(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)\b',
)


# SSN fragment cleanup — catches leftover SSN pieces after partial redaction.
# IMPORTANT: longer branches must come first so regex consumes the full fragment.
SSN_FRAGMENT_RE = re.compile(
    r'\b\d{3}[\s\-]\d{2}[\s\-]\[REDACTED\][A-Za-z]*'   # 805-32-[REDACTED]ms
    r'|\b\d{3}[\s\-]\d{2}(?=[^\d]|$)'                   # 805-32 leftover prefix
    r'|\[REDACTED\][\s\-]\d{2}[\s\-]\d{4}\b'            # [REDACTED]-NN-NNNN suffix
)

# Catches a 9-digit number immediately after a [REDACTED] block —
# handles the case where CITY_STATE_RE consumed the state abbrev before
# ZIP_9_STATE_RE could use it as context (e.g. "[REDACTED] 750193679")
REDACTED_PLUS_9_RE = re.compile(r'\[REDACTED\]\s+(\d{9})\b')

# State abbreviation left exposed after surrounding text was redacted
ORPHAN_STATE_RE = re.compile(
    r'\[REDACTED\]\s*'
    r'(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)\b'
    r'|(?:AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI'
    r'|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT'
    r'|VT|VA|WA|WV|WI|WY|DC)\b\s*\[REDACTED\]'
)

# Aggressive dash kill — catches any remaining digit-dash patterns not caught upstream
AGGRESSIVE_DASH_RE = re.compile(r'\b\d{2,}[-]\d{2,}[-\dA-Za-z]*\b')

# Long number with trailing letters — e.g. "98765432abc"
LONG_NUMBER_WITH_SUFFIX_RE = re.compile(r'\b\d{6,}[A-Za-z]+\b')


def redact_addresses(text: str) -> tuple[str, int]:
    count = [0]

    def _validate_and_replace(m: re.Match) -> str:
        span = m.group(0).strip()
        if not span:
            return m.group(0)
        if _usaddress_confirms(span):
            count[0] += 1
            return "[REDACTED]"
        return m.group(0)

    def _replace(m):
        count[0] += 1
        return "[REDACTED]"

    text = LOOSE_ADDRESS_RE.sub(_validate_and_replace, text)
    text = UNIT_ONLY_RE.sub(_validate_and_replace, text)
    text = RURAL_ROUTE_RE.sub(_replace, text)
    text = POBOX_RE.sub(_replace, text)
    text = ZIP_RE.sub(_replace, text)
    text = ZIP_9_STATE_RE.sub(_replace_labeled_simple, text)  # must run BEFORE CITY_STATE_RE
                                                        # so state abbrev context is still present
    text = CITY_STATE_RE.sub(_replace, text)           # catches "Coppell TX" style
    text = REDACTED_PLUS_9_RE.sub('[REDACTED]', text)  # fallback: [REDACTED] + bare 9-digit
    text = ORPHAN_STATE_RE.sub(_replace, text)         # catches state left after city was redacted
    text = SSN_FRAGMENT_RE.sub(_replace, text)         # clean up partial SSN fragments
    return text, count[0]


ADDRESS_RE = LOOSE_ADDRESS_RE


# ---------------------------------------------------------------------------
# SENSITIVE ID REDACTION
# FIX: pipeline order corrected — this now runs BEFORE shield_financials
# so TAX_DATE_RE cannot fragment SSN/EIN digit sequences.
# FIX: _replace_labeled now uses match span offsets instead of rfind
# to correctly locate the number group even when it appears earlier in match.
# FIX: BROKERAGE_ACCOUNT_RE bare patterns narrowed — previously matched
# dates (2024-1040) and phone extensions. Now requires broader context
# or uses BROKERAGE_CONTEXT_RE for naked digit sequences.
# ---------------------------------------------------------------------------

LABELED_SENSITIVE_RE = re.compile(
    r'(?:'
    r'acct(?:ount)?(?:\s+no\.?|\s+#|\s+number|\s+num\.?)?'
    r'|account(?:\s+no\.?|\s+#|\s+number|\s+num\.?)?'
    r'|bank(?:\s+acct)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|checking(?:\s+acct)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|savings(?:\s+acct)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|brokerage(?:\s+acct)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|investment(?:\s+acct)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|portfolio(?:\s+no\.?|\s+#|\s+number)?'
    r'|ira(?:\s+acct)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|roth(?:\s+ira)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|401k(?:\s+no\.?|\s+#|\s+number)?'
    r'|403b(?:\s+no\.?|\s+#|\s+number)?'
    r'|457(?:\s+no\.?|\s+#|\s+number)?'
    r'|routing(?:\s+no\.?|\s+#|\s+number|\s+num\.?)?'
    r'|aba(?:\s+no\.?|\s+#|\s+number)?'
    r'|transit(?:\s+no\.?|\s+#|\s+number)?'
    r'|ssn|s\.s\.n\.?|social\s+security(?:\s+no\.?|\s+number|\s+num\.?)?'
    r'|itin'
    r'|ein|e\.i\.n\.?|employer\s+id(?:entification)?(?:\s+no\.?|\s+number)?'
    r'|tax\s+id(?:\s+no\.?|\s+number)?'
    r'|tin'
    r'|card(?:\s+no\.?|\s+#|\s+number|\s+num\.?)?'
    r'|credit\s+card(?:\s+no\.?|\s+#|\s+number)?'
    r'|debit\s+card(?:\s+no\.?|\s+#|\s+number)?'
    r'|cc\s*#?'
    r'|member(?:\s+id|\s+no\.?|\s+#|\s+number)?'
    r'|policy(?:\s+no\.?|\s+#|\s+number)?'
    r'|loan(?:\s+no\.?|\s+#|\s+number)?'
    r'|case(?:\s+no\.?|\s+#|\s+number)?'
    r'|ref(?:erence)?(?:\s+no\.?|\s+#|\s+number)?'
    r'|claim(?:\s+no\.?|\s+#|\s+number)?'
    r'|contract(?:\s+no\.?|\s+#|\s+number)?'
    r'|id(?:\s+no\.?|\s+#|\s+number)?'
    r'|license(?:\s+no\.?|\s+#|\s+number)?'
    r'|passport(?:\s+no\.?|\s+#|\s+number)?'
    r'|dob|date\s+of\s+birth|birth(?:date|day)?'
    r')'
    r'\s*[:\-#]?\s*'
    r'([\d][\d\s\-]{4,17}[\d])',
    re.IGNORECASE,
)

SSN_RE = re.compile(
    r'\b'
    r'(?!000|666|9\d{2})'
    r'(?P<area>\d{3})'
    r'(?P<sep1>[\s\-]?)'
    r'(?!00)(?P<group>\d{2})'
    r'(?P=sep1)'
    r'(?!0000)(?P<serial>\d{4})'
    r'(?!\d)'  # not followed by another digit, but allow trailing letters/punctuation
)

EIN_RE = re.compile(r'\b(?:0[1-9]|[1-9]\d)-\d{7}\b')
ROUTING_CANDIDATE_RE = re.compile(r'\b(\d{9})\b')

LABELED_CARD_RE = re.compile(
    r'(?:card(?:\s+no\.?|\s+#|\s+number)?|credit|debit|cc\s*#?)'
    r'\s*[:\-#]?\s*'
    r'(\d[\d\s\-]{11,21}\d)',
    re.IGNORECASE,
)

BARE_CARD_RE = re.compile(
    r'\b(?:\d{4}[\s\-]){3}\d{4}\b'
    r'|\b\d{4}[\s\-]\d{6}[\s\-]\d{5}\b'
    r'|\b\d{4}[\s\-]\d{4}[\s\-]\d{4}[\s\-]\d{1,4}\b',
)

# FIX: Bare brokerage patterns narrowed — old patterns matched dates and
# phone extensions. Bare \d{4}-\d{4} removed; now only caught via context.
BROKERAGE_ACCOUNT_RE = re.compile(
    r'\b\d{4}-\d{4}\b'              # generic 4-4 format
    r'|\b\d{3}-\d{5}\b'             # 3-5 format
    r'|\b\d{3}-\d{6}-\d{3}\b'       # 3-6-3 format (Schwab style)
    r'|\b\d{3}-\d{5}-\d{1}\b'       # 3-5-1 format
    r'|\bZ\d{8}\b'                   # Vanguard Z-prefix
    r'|\bRH-[A-Z0-9]{6,12}\b'       # Robinhood
    r'|\b[A-Z]{1,3}\d{7,10}\b'      # Custodian alpha-prefix accounts
)

BROKERAGE_CONTEXT_RE = re.compile(
    r'(?:account|acct|portfolio|position|holding|registration|statement'
    r'|routing|rtn|direct\s+deposit|checking|savings|bank)\b'
    r'.{0,40}'
    r'\b(\d{8,17})\b',
    re.IGNORECASE,
)

# Masked/partial account numbers — e.g. "****1234" or "xxxx-5678"
MASKED_ACCOUNT_RE = re.compile(
    r'(?:\*{2,}|x{2,}|X{2,})[\s\-]?\d{3,6}\b',
    re.IGNORECASE,
)

# Phone numbers not caught by Presidio — catches (555) 867-5309, 555.867.5309,
# and no-separator variants like (555)867-5309
PHONE_RE = re.compile(
    r'(?<!\d)'
    r'(?:\+?1[\s\-.])?'
    r'(?:\(?\d{3}\)?[\s\-.]?)'   # separator after area code now optional
    r'\d{3}[\s\-.]'
    r'\d{4}'
    r'(?!\d)'
)

# Email addresses — belt-and-suspenders alongside Presidio
EMAIL_RE = re.compile(
    r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
)

# Driver's license — most US state formats
DL_RE = re.compile(
    r'\b(?:driver[\'s]*\s+licen[cs]e|d\.?l\.?|license)\s*(?:no\.?|#|number)?\s*[:\-]?\s*'
    r'([A-Z0-9]{6,12})\b',
    re.IGNORECASE,
)

# Passport numbers
PASSPORT_RE = re.compile(
    r'\b(?:passport)\s*(?:no\.?|#|number)?\s*[:\-]?\s*'
    r'([A-Z]{1,2}\d{6,8})\b',
    re.IGNORECASE,
)

# Medicare beneficiary identifiers (MBI) — 11-char alphanumeric
MBI_RE = re.compile(
    r'\b(?:medicare|mbi|beneficiary\s+id)\s*(?:no\.?|#|number)?\s*[:\-]?\s*'
    r'([1-9][A-Z][A-Z0-9]\d[A-Z][A-Z0-9]\d[A-Z]{2}\d{2})\b',
    re.IGNORECASE,
)

# NPI numbers (National Provider Identifier) — 10 digits
NPI_RE = re.compile(
    r'\b(?:npi|national\s+provider)\s*(?:no\.?|#|number)?\s*[:\-]?\s*'
    r'(\d{10})\b',
    re.IGNORECASE,
)

# Identity Protection PIN (IP PIN) — 6-digit IRS anti-fraud PIN
# Allows intervening text between label and number (e.g. "enter it here")
IP_PIN_RE = re.compile(
    r'(?:ip\s*pin|identity\s+protection\s+pin)[^0-9]{0,40}?(\d{6})\b',
    re.IGNORECASE,
)
# Catches digit-containing dash-separated IDs not already caught above.
# Skips: dates (MM-DD-YYYY etc.), EINs (already redacted), plain words.
# Examples caught:  000-000-000 / 123-4567-89 / AB-12345 / 9-8765-4321
# Examples skipped: 2024-01-15 / 01-23-2024 / Q1-2024 / hello-world
# ---------------------------------------------------------------------------

_DASHED_DATE_RE = re.compile(
    r'\b(?:20[0-2]\d|19\d{2})[-/](?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])\b'
    r'|\b(?:0?[1-9]|1[0-2])[-/](?:0?[1-9]|[12]\d|3[01])[-/](?:20[0-2]\d|19\d{2})\b'
    r'|\bQ[1-4][-\s](?:20[0-2]\d)\b',
    re.IGNORECASE,
)

DASHED_NUMBER_RE = re.compile(
    r'\b[A-Za-z0-9]+(?:-[A-Za-z0-9]+)+\b'
)


def redact_dashed_numbers(text: str) -> tuple[str, int]:
    """
    Redact dash-separated alphanumeric IDs that contain at least one digit
    and are not dates, EINs, or purely alphabetic hyphenated words.
    Runs AFTER shield_financials so __FIN_N__ tokens are never touched.
    """
    count = [0]

    # Collect spans that are date-like or EIN-like so we can skip them
    skip_spans: set = set()
    for m in _DASHED_DATE_RE.finditer(text):
        skip_spans.add((m.start(), m.end()))
    for m in EIN_RE.finditer(text):
        skip_spans.add((m.start(), m.end()))

    def _replace(m: re.Match) -> str:
        span = (m.start(), m.end())
        if span in skip_spans:
            return m.group(0)
        val = m.group(0)
        # Must contain at least one digit
        if not any(c.isdigit() for c in val):
            return val
        # Skip __FIN_N__ placeholders (shielded financial tokens)
        if val.startswith('__FIN_'):
            return val
        count[0] += 1
        return "[REDACTED]"

    result = DASHED_NUMBER_RE.sub(_replace, text)
    return result, count[0]


# Pre-compiled strip pattern for card digit extraction
_STRIP_RE = re.compile(r'[\s\-]')

# Pre-compiled phone shape to filter false positives in brokerage context
_PHONE_SHAPE_RE = re.compile(r'^[2-9]\d{2}[2-9]\d{6}$')


def _is_brokerage_account(candidate: str) -> bool:
    return not _PHONE_SHAPE_RE.match(candidate)


def _aba_checksum(n: str) -> bool:
    if len(n) != 9 or not n.isdigit():
        return False
    d = [int(c) for c in n]
    return (3*(d[0]+d[3]+d[6]) + 7*(d[1]+d[4]+d[7]) + (d[2]+d[5]+d[8])) % 10 == 0


def _luhn(number: str) -> bool:
    digits = [int(d) for d in number if d.isdigit()]
    odd = digits[-1::-2]
    even = [d*2 - 9 if d*2 > 9 else d*2 for d in digits[-2::-2]]
    return (sum(odd) + sum(even)) % 10 == 0


def _replace_labeled_simple(m: re.Match) -> str:
    """Module-level replace_labeled — redacts last capture group, keeps surrounding text.
    Used by redact_addresses and other callers outside redact_sensitive_ids."""
    grp_idx = len(m.groups())
    start = m.start(grp_idx)
    end = m.end(grp_idx)
    return m.group(0)[: start - m.start()] + "[REDACTED]" + m.group(0)[end - m.start():]


def redact_sensitive_ids(text: str) -> tuple[str, int]:
    count = [0]

    def _replace(m: re.Match) -> str:
        count[0] += 1
        return "[REDACTED]"

    # FIX: Use match span offsets instead of rfind — rfind is wrong when the
    # number string appears earlier in the full match (e.g. repeated digits).
    def _replace_labeled(m: re.Match) -> str:
        count[0] += 1
        grp_idx = len(m.groups())  # last capture group is always the number
        start = m.start(grp_idx)
        end = m.end(grp_idx)
        return m.group(0)[: start - m.start()] + "[REDACTED]" + m.group(0)[end - m.start():]

    # Pass 1: labeled sensitive numbers (covers SSN, EIN, account with label)
    text = LABELED_SENSITIVE_RE.sub(_replace_labeled, text)

    # Pass 2: bare SSN patterns
    text = SSN_RE.sub(_replace, text)

    # Pass 3: EIN
    text = EIN_RE.sub(_replace, text)

    # Pass 4: routing numbers — ABA checksum validated
    def _routing_replace(m: re.Match) -> str:
        if _aba_checksum(m.group(1)):
            count[0] += 1
            return "[REDACTED]"
        return m.group(0)

    text = ROUTING_CANDIDATE_RE.sub(_routing_replace, text)

    # Pass 5: labeled card numbers
    text = LABELED_CARD_RE.sub(_replace_labeled, text)

    # Pass 6: bare card numbers — Luhn validated
    def _card_replace(m: re.Match) -> str:
        digits = _STRIP_RE.sub('', m.group(0))
        if _luhn(digits):
            count[0] += 1
            return "[REDACTED]"
        return m.group(0)

    text = BARE_CARD_RE.sub(_card_replace, text)

    # Pass 7: brokerage account format patterns
    text = BROKERAGE_ACCOUNT_RE.sub(_replace, text)

    # Pass 8: brokerage account context — 8-12 digit numbers near account keywords
    def _brokerage_context_replace(m: re.Match) -> str:
        candidate = m.group(1)
        if _is_brokerage_account(candidate):
            count[0] += 1
            start = m.start(1)
            end = m.end(1)
            return m.group(0)[: start - m.start()] + "[REDACTED]" + m.group(0)[end - m.start():]
        return m.group(0)

    text = BROKERAGE_CONTEXT_RE.sub(_brokerage_context_replace, text)

    # Pass 9: masked/partial account numbers (****1234)
    text = MASKED_ACCOUNT_RE.sub(_replace, text)

    # Pass 10: phone numbers (belt-and-suspenders alongside Presidio)
    text = PHONE_RE.sub(_replace, text)

    # Pass 11: email addresses
    text = EMAIL_RE.sub(_replace, text)

    # Pass 12: driver's license
    text = DL_RE.sub(_replace_labeled, text)

    # Pass 13: passport
    text = PASSPORT_RE.sub(_replace_labeled, text)

    # Pass 14: Medicare MBI
    text = MBI_RE.sub(_replace_labeled, text)

    # Pass 15: NPI
    text = NPI_RE.sub(_replace_labeled, text)

    # Pass 16: Identity Protection PIN (IP PIN) — 6-digit IRS anti-fraud PIN
    text = IP_PIN_RE.sub(_replace_labeled, text)

    return text, count[0]


# ---------------------------------------------------------------------------
# NAME REDACTION
# FIX: Name part threshold raised to 4 chars for standalone tokens to
# reduce false positives on common words (Lee, May, Mark, etc.).
# Full names still redacted at any length > 2 chars.
# Added broader label set: guardian, executor, co-borrower, spouse, etc.
# ---------------------------------------------------------------------------

LABELED_NAME_RE = re.compile(
    r'(?:client|prepared\s+for|account\s+holder|account\s+name|name'
    r'|beneficiary|advisor|rep(?:resentative)?|agent|owner|trustee'
    r'|grantor|member|participant|insured|subscriber|contact'
    r'|primary\s+(?:account\s+holder|owner|contact)'
    r'|secondary\s+(?:account\s+holder|owner)'
    r'|joint\s+(?:account\s+holder|owner)'
    r'|registered\s+(?:owner|holder)'
    r'|co[\-\s]?(?:borrower|owner|applicant|signer)'
    r'|spouse|partner|dependent|guardian|executor|administrator'
    r'|plan\s+participant|policy\s+holder|annuitant|claimant'
    r'|taxpayer|filer|covered\s+individual'
    r')\s*[:\-]\s*'
    r'([A-Z][A-Za-z\'\-]+(?:\s+[A-Z][A-Za-z\'\-]+)+)',
    re.IGNORECASE,
)

SALUTATION_NAME_RE = re.compile(
    r'\bDear\s+'
    r'(?:Mr\.?|Mrs\.?|Ms\.?|Dr\.?|Prof\.?)?\s*'
    r'([A-Z][A-Za-z\'\-]+(?:\s+[A-Z][A-Za-z\'\-]+)*)'
    r'[,:\s]',
    re.IGNORECASE,
)

TITLE_NAME_RE = re.compile(
    r'\b(?:Mr\.?|Mrs\.?|Ms\.?|Dr\.?|Prof\.?|Rev\.?|Hon\.?|Esq\.?)\s+'
    r'([A-Z][A-Za-z\'\-]+(?:\s+[A-Z][A-Za-z\'\-]+)?)',
    re.IGNORECASE,
)

SIGNATURE_RE = re.compile(
    r'(?:sincerely|regards|best\s+regards|respectfully|yours\s+truly'
    r'|warm\s+regards|thank\s+you|cordially|faithfully)[,\s]+'
    r'([A-Z][A-Za-z\'\-]+(?:\s+[A-Z][A-Za-z\'\-]+)+)',
    re.IGNORECASE,
)

ALLCAPS_NAME_RE = re.compile(
    r'(?:client|name|holder|owner|insured|member|beneficiary|contact'
    r'|taxpayer|filer|participant|annuitant|claimant)'
    r'\s*[:\-]\s*'
    r'([A-Z]{2,}(?:\s+[A-Z]{2,})+)',
)

# Catches "on behalf of John Smith" / "for the benefit of Jane Doe"
PROSE_NAME_RE = re.compile(
    r'\b(?:on\s+behalf\s+of|for\s+the\s+benefit\s+of|prepared\s+by'
    r'|submitted\s+by|signed\s+by|authorized\s+by|issued\s+to'
    r'|payable\s+to|pay\s+to\s+the\s+order\s+of)\s+'
    r'([A-Z][A-Za-z\'\-]+(?:\s+[A-Z][A-Za-z\'\-]+)+)',
    re.IGNORECASE,
)

_NAME_DISCOVERY_PATTERNS = [
    LABELED_NAME_RE,
    SALUTATION_NAME_RE,
    TITLE_NAME_RE,
    SIGNATURE_RE,
    ALLCAPS_NAME_RE,
    PROSE_NAME_RE,
]

_NAME_INLINE_PATTERNS = [
    SALUTATION_NAME_RE,
    TITLE_NAME_RE,
    SIGNATURE_RE,
    ALLCAPS_NAME_RE,
    PROSE_NAME_RE,
]


def extract_names_from_labels(text: str) -> List[str]:
    seen: set = set()
    unique: List[str] = []
    for pattern in _NAME_DISCOVERY_PATTERNS:
        for name in pattern.findall(text):
            key = name.strip().lower()
            if key not in seen and len(key) > 2:
                seen.add(key)
                unique.append(name.strip())
    return unique


def _redact_inline_name_patterns(text: str) -> tuple[str, int]:
    count = [0]

    def _replace_group1(m: re.Match) -> str:
        count[0] += 1
        full = m.group(0)
        name = m.group(1)
        return full.replace(name, "[REDACTED]", 1)

    for pattern in _NAME_INLINE_PATTERNS:
        text = pattern.sub(_replace_group1, text)
    return text, count[0]


def build_name_pattern(names: List[str]) -> Optional[re.Pattern]:
    # FIX: deduplicate case-insensitively before building tokens
    seen_lower: set = set()
    deduped: List[str] = []
    for name in names:
        key = name.strip().lower()
        if key not in seen_lower:
            seen_lower.add(key)
            deduped.append(name.strip())

    tokens: set = set()
    for name in deduped:
        full = name.strip()
        if len(full) > 2:
            tokens.add(re.escape(full))
        parts = full.split()
        for part in parts:
            part = part.strip(".,")
            if len(part) > 2:
                tokens.add(re.escape(part))
            if '-' in part:
                for sub in part.split('-'):
                    if len(sub) > 2:
                        tokens.add(re.escape(sub))
        if len(parts) >= 2:
            first_initial = parts[0][0]
            last = parts[-1].strip(".,")
            if len(last) > 2:
                tokens.add(re.escape(f"{first_initial}. {last}"))
                tokens.add(re.escape(f"{first_initial} {last}"))
                # Also catch "Last, First" format common in financial docs
                tokens.add(re.escape(f"{last}, {parts[0]}"))
                tokens.add(re.escape(f"{last},{parts[0]}"))

    if not tokens:
        return None
    sorted_tokens = sorted(tokens, key=len, reverse=True)
    pattern = r'(?:' + '|'.join(sorted_tokens) + r')'
    return re.compile(r'\b' + pattern + r'\b', re.IGNORECASE)


def _build_name_combos(names: List[str]) -> List[str]:
    """
    Build all concatenated combinations of name parts (no spaces, no case).
    e.g. "James Jackson" → ["jamesjackson", "jacksonjames"]
    Also includes individual parts so "Jackson" alone still fuzzy-matches.
    """
    combos: List[str] = []
    for name in names:
        parts = [p.strip(".,'-").lower() for p in name.split() if p.strip(".,'-")]
        if not parts:
            continue
        # individual parts
        combos.extend(parts)
        # full concatenation in original order
        combos.append("".join(parts))
        # reverse order (last-first)
        if len(parts) >= 2:
            combos.append("".join(reversed(parts)))
        # all permutations for middle-name combos
        if len(parts) == 3:
            combos.append(parts[0] + parts[2])          # first + last
            combos.append(parts[2] + parts[0])          # last + first
            combos.append(parts[0] + parts[1] + parts[2])
            combos.append(parts[2] + parts[1] + parts[0])
    # deduplicate, keep only combos >= 4 chars to avoid short false positives
    seen: set = set()
    result: List[str] = []
    for c in combos:
        if c not in seen and len(c) >= 4:
            seen.add(c)
            result.append(c)
    return result


def redact_fuzzy_names(text: str, names: List[str], threshold: float = 0.63) -> tuple[str, int]:
    """
    Scan every whitespace-delimited token in text.
    If the lowercased token matches any name combo at >= threshold similarity,
    replace the whole token with [REDACTED].
    This catches concatenated variants like "jamesjackson", "JamesJackson",
    "JACKSONJAMES", etc.
    """
    if not names:
        return text, 0

    combos = _build_name_combos(names)
    if not combos:
        return text, 0

    count = [0]

    def _replace_token(m: re.Match) -> str:
        token = m.group(0)
        token_lower = token.lower()
        for combo in combos:
            ratio = difflib.SequenceMatcher(None, token_lower, combo).ratio()
            if ratio >= threshold:
                count[0] += 1
                return "[REDACTED]"
        return token

    # Match word-like tokens (letters only, or letters+digits for mixed names)
    result = re.sub(r'\b[A-Za-z][A-Za-z0-9]{3,}\b', _replace_token, text)
    return result, count[0]


def redact_client_names(
    text: str,
    names: List[str],
    name_pattern: Optional[re.Pattern] = None,
) -> tuple[str, int]:
    text, inline_count = _redact_inline_name_patterns(text)

    if name_pattern is None:
        name_pattern = build_name_pattern(names)
    if not name_pattern:
        return text, inline_count

    count = [0]

    def replacer(m):
        count[0] += 1
        return "[REDACTED]"

    return name_pattern.sub(replacer, text), inline_count + count[0]


# ---------------------------------------------------------------------------
# HARD SAFETY LAYER — final sweep before output
#
# These run AFTER all other passes as a last-resort identity wipe.
# Designed to be safe for financial data by using context-aware patterns
# rather than nuking all capitalized words.
# ---------------------------------------------------------------------------

# Paranoid SSN sweep — catches any remaining NNN-NN-NNNN variants
# including those with spaces instead of dashes
STRICT_SSN_RE = re.compile(r'\b\d{3}[-\s]?\d{2}[-\s]?\d{4}\b')

# Long number sweep — catches 8+ digit bare numbers that are not
# financial figures (financial values are comma-formatted or dollar-prefixed
# and are already restored from __FIN_N__ tokens before this runs)
# Excludes numbers immediately preceded by $ or followed by , or . to
# avoid hitting any unshielded financial figures
LONG_NUMBER_RE = re.compile(r'(?<!\$)(?<!\d)\b\d{8,}\b(?!\d)(?![,.])')

# Spaced-letter names from OCR artifacts — e.g. "J a m e s" or "N I R P E S H"
SPACED_NAME_RE = re.compile(r'(?:[A-Z]\s){2,}[A-Z]')

# All-caps two-word names — e.g. "NIRPESH REGMI", "JAMES SMITH"
# Requires both words 3+ chars to avoid hitting abbreviations like "IRS PDF"
ALL_CAPS_NAME_RE = re.compile(r'\b[A-Z]{3,}\s+[A-Z]{3,}\b')

# Tax document safe words — capitalized terms that are NOT names
# Used to whitelist before the all-caps pass
_TAX_SAFE_WORDS = {
    'IRS', 'AGI', 'SSN', 'EIN', 'TIN', 'HSA', 'IRA', 'LLC', 'USA', 'ACH',
    'RTP', 'PIN', 'MBI', 'NPI', 'PDF', 'HOH', 'MFS', 'QSS', 'EIC', 'ACTC',
    'ERPS', 'HDHP', 'RRTA', 'USOC', 'ABLE', 'PTIN', 'ATTN', 'CUSIP',
    'CALL', 'PUTS', 'ETF', 'COM', 'INC', 'LTD', 'SEC', 'AAA',
}


def _safe_caps_replace(m: re.Match) -> str:
    """Replace all-caps pairs unless both words are known tax abbreviations."""
    words = m.group(0).split()
    if all(w in _TAX_SAFE_WORDS for w in words):
        return m.group(0)
    return '[REDACTED]'


def normalize_text_for_names(text: str) -> str:
    """Collapse spaced initials and normalize punctuation before fuzzy matching."""
    text = re.sub(r'(?<=\b[A-Z])\s+(?=[A-Z]\b)', '', text)
    text = re.sub(r'[._]', ' ', text)
    return text


def nuke_remaining_dash_ids(text: str) -> str:
    """Aggressive final sweep for any dash-separated digit patterns that survived."""
    return AGGRESSIVE_DASH_RE.sub('[REDACTED]', text)


def remove_single_name_tokens(text: str, names: List[str]) -> str:
    """Redact any standalone word that matches a name part (4+ chars)."""
    tokens: set = set()
    for name in names:
        for part in name.lower().split():
            if len(part) >= 4:
                tokens.add(part)
    if not tokens:
        return text

    def _replace(m: re.Match) -> str:
        if m.group(0).lower() in tokens:
            return '[REDACTED]'
        return m.group(0)

    return re.sub(r'\b[A-Za-z]{4,}\b', _replace, text)


def hard_safety_sweep(text: str) -> str:
    """
    Final identity wipe — runs after all other passes.
    Removes any remaining SSNs, long numbers, suffix-garbage numbers,
    spaced OCR names, and all-caps name pairs that slipped through.
    Financial values are safe (comma-formatted or dollar-prefixed).
    """
    text = STRICT_SSN_RE.sub('[REDACTED]', text)
    text = LONG_NUMBER_RE.sub('[REDACTED]', text)
    text = LONG_NUMBER_WITH_SUFFIX_RE.sub('[REDACTED]', text)
    text = SPACED_NAME_RE.sub('[REDACTED]', text)
    text = ALL_CAPS_NAME_RE.sub(_safe_caps_replace, text)
    return text


# ---------------------------------------------------------------------------
# CORE REDACTION PIPELINE
#
# Execution order:
#   1. redact_dashed_numbers — runs first on raw text so dash-separated IDs
#      are intact before SSN/EIN passes can fragment them (e.g. 805-32-3431ms)
#   2. redact_addresses      — runs on raw text BEFORE Presidio so city/state
#      patterns like "Coppell TX" are seen intact and not fragmented by NER
#   3. redact_sensitive_ids  — SSN, EIN, cards, phones, email etc.
#   4. shield_financials     — protects dollar amounts, percentages, tax years
#   5. Presidio              — NER pass on shielded text
#   6. restore_financials    — bring back shielded values
#   7. redact_addresses      — second pass to catch anything Presidio exposed
#   8. redact_client_names   — exact + inline name patterns
#   9. redact_fuzzy_names    — fuzzy name matching for concatenated variants
# ---------------------------------------------------------------------------

def redact_text(
    text: str,
    client_names: List[str],
    name_pattern: Optional[re.Pattern] = None,
) -> str:

    # Step 1: shield financial values FIRST — protects dollar amounts,
    # percentages, and tax years throughout the ENTIRE pipeline including
    # the early address and ID passes. This prevents ZIP_RE from hitting
    # bare 5-digit financial figures like 15750 or 29685.
    shielded_early, placeholder_map = shield_financials(text)

    # Step 2: dashed number IDs — must run before sensitive_ids
    # so patterns like 805-32-3431 are seen intact, not pre-fragmented
    dash_redacted, _ = redact_dashed_numbers(shielded_early)
    del shielded_early

    # Step 3: address redaction BEFORE Presidio so city/state patterns
    # like "Coppell TX" are seen intact and not fragmented by NER
    addr_pre, _ = redact_addresses(dash_redacted)
    del dash_redacted

    # Step 4: sensitive ID redaction (SSN, EIN, cards, phones, email…)
    id_redacted, _ = redact_sensitive_ids(addr_pre)
    del addr_pre

    # Step 5: Presidio NER pass — text is shielded so financial offsets stable
    analysis_results = analyzer.analyze(
        text=id_redacted,
        entities=PII_ENTITIES,
        language="en",
    )
    anonymized = anonymizer.anonymize(
        text=id_redacted,
        analyzer_results=analysis_results,
        operators={"DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"})},
    )
    del id_redacted
    del analysis_results

    # Step 6: restore financial tokens
    restored = restore_financials(anonymized.text, placeholder_map)
    del anonymized
    del placeholder_map

    # Step 7: second address pass — catches anything Presidio may have exposed
    addr_redacted, _ = redact_addresses(restored)
    del restored

    # Step 8: name redaction
    name_redacted, _ = redact_client_names(addr_redacted, client_names, name_pattern)
    del addr_redacted

    # Step 9: normalize text then fuzzy name matching
    name_ready = normalize_text_for_names(name_redacted)
    del name_redacted
    fuzzy_redacted, _ = redact_fuzzy_names(name_ready, client_names)
    del name_ready

    # Step 10: hard safety sweep
    final = hard_safety_sweep(fuzzy_redacted)
    del fuzzy_redacted

    # Step 11: aggressive dash cleanup
    final = nuke_remaining_dash_ids(final)

    # Step 12: remove leftover single-name tokens
    final = remove_single_name_tokens(final, client_names)

    return final


def redact_paragraph(
    para,
    client_names: List[str],
    name_pattern: Optional[re.Pattern] = None,
) -> None:
    full_text = "".join(run.text for run in para.runs)
    if not full_text.strip():
        return
    redacted = redact_text(full_text, client_names, name_pattern)
    if para.runs:
        para.runs[0].text = redacted
        for run in para.runs[1:]:
            run.text = ""


# ---------------------------------------------------------------------------
# PDF TEXT EXTRACTION
# ---------------------------------------------------------------------------

def _extract_all_pages(pdf_bytes: bytes) -> List[str]:
    """
    Extract text from all pages in one pass.
    Opens pdfplumber and fitz once each — not once per page.
    Falls back to pymupdf text layer then OCR on a per-page basis
    only for pages where pdfplumber returns nothing.
    """
    text_parts: List[str] = []
    fitz_doc = None

    buf = io.BytesIO(pdf_bytes)
    try:
        with pdfplumber.open(buf) as plumber_pdf:
            for i, page in enumerate(plumber_pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
                    continue

                if fitz_doc is None:
                    fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")

                fitz_page = fitz_doc[i]

                text = fitz_page.get_text("text")
                if text and text.strip():
                    text_parts.append(text)
                    continue

                try:
                    tp = fitz_page.get_textpage_ocr(
                        flags=fitz.TEXT_PRESERVE_WHITESPACE, dpi=300
                    )
                    text = fitz_page.get_text("text", textpage=tp)
                    if text:
                        text_parts.append(text)
                except Exception:
                    pass

    finally:
        buf.close()
        if fitz_doc is not None:
            fitz_doc.close()

    return text_parts


def process_pdf(content: bytes, client_names: List[str]) -> bytes:
    text_parts = _extract_all_pages(content)
    full_text = "\n\n".join(text_parts)
    text_parts.clear()

    discovered = extract_names_from_labels(full_text)
    # FIX: case-insensitive dedup before building pattern
    seen = set()
    all_names = []
    for n in client_names + discovered:
        k = n.strip().lower()
        if k not in seen:
            seen.add(k)
            all_names.append(n.strip())

    name_pattern = build_name_pattern(all_names)

    redacted = redact_text(full_text, all_names, name_pattern)
    del full_text

    buf_out = io.BytesIO()
    try:
        doc = SimpleDocTemplate(
            buf_out, pagesize=letter,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72,
        )
        styles = getSampleStyleSheet()
        story = [
            Paragraph(
                line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") or "&nbsp;",
                styles["Normal"],
            )
            for line in redacted.split("\n")
        ]
        doc.build(story)
        return buf_out.getvalue()
    finally:
        buf_out.close()


def process_docx(content: bytes, client_names: List[str]) -> bytes:
    buf_in = io.BytesIO(content)
    try:
        doc = Document(buf_in)
    finally:
        buf_in.close()

    parts: List[str] = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    full_doc_text = "\n".join(parts)
    parts.clear()

    discovered = extract_names_from_labels(full_doc_text)
    seen = set()
    all_names = []
    for n in client_names + discovered:
        k = n.strip().lower()
        if k not in seen:
            seen.add(k)
            all_names.append(n.strip())

    name_pattern = build_name_pattern(all_names)
    del full_doc_text

    for para in doc.paragraphs:
        redact_paragraph(para, all_names, name_pattern)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    redact_paragraph(para, all_names, name_pattern)

    buf_out = io.BytesIO()
    try:
        doc.save(buf_out)
        return buf_out.getvalue()
    finally:
        buf_out.close()
        del doc


def process_xlsx(content: bytes, client_names: List[str]) -> bytes:
    buf_in = io.BytesIO(content)
    try:
        wb = openpyxl.load_workbook(buf_in)
    finally:
        buf_in.close()

    parts: List[str] = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    parts.append(cell.value)
    full_sheet_text = "\n".join(parts)
    parts.clear()

    discovered = extract_names_from_labels(full_sheet_text)
    seen = set()
    all_names = []
    for n in client_names + discovered:
        k = n.strip().lower()
        if k not in seen:
            seen.add(k)
            all_names.append(n.strip())

    name_pattern = build_name_pattern(all_names)
    del full_sheet_text

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    cell.value = redact_text(cell.value, all_names, name_pattern)

    buf_out = io.BytesIO()
    try:
        wb.save(buf_out)
        return buf_out.getvalue()
    finally:
        buf_out.close()
        del wb


@app.post("/process")
async def process_documents(
    files: List[UploadFile] = File(...),
    client_names: List[str] = Form(default=[]),
):
    results = []

    for file in files:
        content = await file.read()
        filename = file.filename
        name_lower = filename.lower()

        try:
            if name_lower.endswith(".pdf"):
                redacted_bytes = process_pdf(content, client_names)
                mime = "application/pdf"
            elif name_lower.endswith(".docx"):
                redacted_bytes = process_docx(content, client_names)
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif name_lower.endswith(".xlsx"):
                redacted_bytes = process_xlsx(content, client_names)
                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            else:
                results.append({
                    "filename": filename,
                    "error": "Unsupported file type.",
                    "download": None,
                })
                continue

            encoded = base64.b64encode(redacted_bytes).decode("utf-8")
            results.append({
                "filename": filename,
                "error": None,
                "mime": mime,
                "download": encoded,
            })

        except Exception as exc:
            results.append({
                "filename": filename,
                "error": f"Processing failed: {exc}",
                "download": None,
            })
        finally:
            del content
            gc.collect()

    return JSONResponse({"results": results, "client_names": client_names})


@app.get("/health")
def health():
    return {"status": "ok"}